import json
import psycopg2
import os
import boto3
from datetime import datetime
from io import BytesIO
from docx import Document

# AWS S3 Configuration
S3_BUCKET = "karini-demo-data-us-east-1"
S3_FOLDER = "Demand-forecasting/order-confirmation/"

# AWS Lambda handler
def lambda_handler(event, context):
    try:
        # Extract parameters from event
        item_id = event.get('item_id')
        location = event.get('location')
        y = event.get('y')
        product_category = event.get('product_category')

        # Validate input
        if not item_id or not location or not y or not product_category:
            return {
                'statusCode': 400,
                'body': json.dumps("Missing required fields")
            }

        # Auto-generate timestamp (ds)
        ds = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')

        # Fetch database credentials from environment variables
        db_host = os.getenv('DB_HOST')
        db_port = os.getenv('DB_PORT', '5432')
        db_name = os.getenv('DB_NAME')
        db_user = os.getenv('DB_USER')
        db_password = os.getenv('DB_PASSWORD')

        # Connect to PostgreSQL
        conn = psycopg2.connect(
        #    dbname=db_name,
            user=db_user,
            password=db_password,
            host=db_host,
            port=db_port
        )
        cursor = conn.cursor()

        # Insert data into the database
        insert_query = """
        INSERT INTO synthetic_demand_data_table (item_id, location, ds, y, product_category)
        VALUES (%s, %s, %s, %s, %s)
        """
        cursor.execute(insert_query, (item_id, location, ds, y, product_category))
        conn.commit()

        # Close database connection
        cursor.close()
        conn.close()

        # Generate DOCX order confirmation
        docx_content = generate_order_confirmation_docx(item_id, location, ds, y, product_category)

        # Upload DOCX to S3
        file_name = f"order_confirmation_{timestamp}.docx"
        s3_path = f"{S3_FOLDER}{file_name}"
        upload_to_s3(docx_content, S3_BUCKET, s3_path)

        return {
            'statusCode': 200,
            'body': json.dumps({
                "message": "Order Confirmation Data Inserted Successfully!",
                "s3_file_path": f"s3://{S3_BUCKET}/{s3_path}",
                "download_link": f"s3://{S3_BUCKET}/{s3_path}"
            })
        }
    
    except Exception as e:
        return {
            'statusCode': 500,
            'body': json.dumps(f"Error: {str(e)}")
        }

# Function to generate DOCX order confirmation
def generate_order_confirmation_docx(item_id, location, ds, y, product_category):
    buffer = BytesIO()
    doc = Document()

    # Title
    doc.add_heading('Order Confirmation', level=1)

    # Order details
    doc.add_paragraph(f"**Order Date:** {ds}")
    doc.add_paragraph(f"**Item ID:** {item_id}")
    doc.add_paragraph(f"**Location:** {location}")
    doc.add_paragraph(f"**Quantity (y):** {y}")
    doc.add_paragraph(f"**Product Category:** {product_category}")

    # Footer
    doc.add_paragraph("\nThank you for your order!", style='Intense Quote')

    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to upload DOCX to S3
def upload_to_s3(docx_content, bucket, key):
    s3_client = boto3.client('s3')
    s3_client.upload_fileobj(docx_content, bucket, key)
